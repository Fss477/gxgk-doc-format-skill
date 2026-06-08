#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
广西国控集团公文格式排版工具
根据广西国控资本运营集团有限责任公司的公文格式规范自动排版 Word 文档
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 字体定义
FONT_TITLE = "方正小标宋简体"
FONT_BODY = "仿宋_GB2312"
FONT_HEADING1 = "黑体"
FONT_HEADING2 = "楷体_GB2312"
FONT_FOOTER = "宋体"

# 字号定义（pt）
SIZE_TITLE = 22  # 二号
SIZE_BODY = 16   # 三号
SIZE_TABLE_TITLE = 18  # 小二
SIZE_TABLE_BODY = 14   # 小四
SIZE_CONTACT = 14      # 小四

# 段落格式
LINE_SPACING_BODY = Pt(29)    # 正文固定行距 29pt
LINE_SPACING_HEADING = Pt(33) # 标题固定行距 33pt
FIRST_LINE_INDENT = Pt(29)    # 首行缩进 2 字符（约 29pt）


def set_paragraph_format(paragraph, font_name=FONT_BODY, font_size=SIZE_BODY, 
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                         first_line_indent=True, bold=False, is_heading=False):
    """设置段落格式"""
    # 设置对齐方式
    paragraph.alignment = alignment
    
    # 设置行距（标题使用 33pt，正文使用 29pt）
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = LINE_SPACING_HEADING if is_heading else LINE_SPACING_BODY
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    # 设置首行缩进
    if first_line_indent and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        paragraph_format.first_line_indent = FIRST_LINE_INDENT
    else:
        paragraph_format.first_line_indent = Pt(0)
    
    # 设置字体
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def detect_heading_level(text):
    """检测标题层级"""
    if not text:
        return 0
    
    # 一级标题：一、二、三、
    if re.match(r'^[一二三四五六七八九十]+[、,.]', text):
        return 1
    
    # 二级标题：（一）（二）
    if re.match(r'^[（(][一二三四五六七八九十]+[）)][、,.]?', text):
        return 2
    
    # 三级标题：1. 1、但需要排除正文中的列举项
    # 判断标准：如果段落较短（<30 字符）且以数字开头，可能是标题
    if re.match(r'^\d+[.,]', text):
        if len(text) < 30:  # 短段落，可能是标题
            return 3
        # 长段落，可能是正文中的列举，不作为标题
    
    # 四级标题：（1）(1)
    if re.match(r'^[（(]\d+[）)]', text):
        return 4
    
    return 0


def is_short_heading(text):
    """判断是否为短标题（如"工作目标"、"重点任务"等）"""
    if not text:
        return False
    # 短标题特征：2-10 个字符，没有标点符号结尾
    if 2 <= len(text) <= 10:
        if not text.endswith(('.', ',', ',', '!', '?', ':')):
            # 常见的标题关键词
            heading_keywords = [
                '目标', '任务', '原则', '要求', '重点', '措施', '步骤', '安排', 
                '组织', '保障', '实施', '背景', '意义', '思路', '情况', 
                '必要性', '原因', '依据', '范围', '对象', '条件', '内容',
                '方式', '计划', '方案', '建议', '意见', '总结', '报告', '请示',
                '预算', '制度', '机制', '体系', '平台', '建设'
            ]
            for kw in heading_keywords:
                if kw in text:
                    return True
    return False


def is_date_line(text):
    """判断是否为日期行"""
    # 匹配格式：2026 年 5 月  日 或 2026 年 5 月 19 日
    if re.match(r'^\d{4}年\d{1,2}月\d{0,2}日', text):
        return True
    if re.match(r'^（\d{4}年\d{1,2}月\d{1,2}日）', text):
        return True
    return False


def is_main_subject(text):
    """判断是否为主送单位"""
    patterns = [
        r'^各部门、',
        r'^各(.*)公司[：:]',
        r'^各(.*)单位 [：:]',
        r'^公司各部门：',
        r'^集团各部门:',
    ]
    for pattern in patterns:
        if re.match(pattern, text):
            return True
    return False


def is_attachment(text):
    """判断是否为附件说明"""
    return text.startswith('附件') or text.startswith('附件 1')


def is_sender(text):
    """判断是否为发文单位"""
    # 发文单位通常是短文本，且不含标点符号
    if len(text) < 15 and not text.endswith(('.', ',', ',', '!')):
        if '公司' in text or '集团' in text or '办公室' in text or '部门' in text:
            return True
    return False


def add_heading_number(paragraph, level, counter):
    """为标题添加序号"""
    if not paragraph.runs:
        return
    
    # 中文数字映射表（避免使用 chr() 导致的编码问题）
    chinese_numerals = '零一二三四五六七八九十十一十二十三十四十五十六十七十八十九二十二十一二十二二十三二十四二十五二十六二十七二十八二十九三十三一三二三三三四三五三六三七三八三九四十四一四二四三四四四五四六四七四八四九五十'
    
    # 根据层级生成序号
    if level == 1:
        # 一级标题：一、二、三、
        counter[0] += 1
        if counter[0] <= 50:
            num_text = chinese_numerals[counter[0]] + '、'
        else:
            num_text = f'第{counter[0]}部分、'
    elif level == 2:
        # 二级标题：（一）（二）
        counter[1] += 1
        if counter[1] <= 50:
            num_text = '(' + chinese_numerals[counter[1]] + ')'
        else:
            num_text = f'(第{counter[1]}部分)'
    elif level == 3:
        # 三级标题：1. 2. 3.
        counter[2] += 1
        num_text = str(counter[2]) + '.'
    elif level == 4:
        # 四级标题：（1）（2）
        counter[3] += 1
        num_text = '(' + str(counter[3]) + ')'
    else:
        return
    
    # 检查是否已有相同层级的序号
    text = paragraph.text.strip()
    has_number = False
    if level == 1 and re.match(r'^[一二三四五六七八九十]+[、,.]', text):
        has_number = True
    elif level == 2 and re.match(r'^[（(][一二三四五六七八九十]+[）)]', text):
        has_number = True
    elif level == 3 and re.match(r'^\d+[.,]', text):
        has_number = True
    elif level == 4 and re.match(r'^[（(]\d+[）)]', text):
        has_number = True
    
    # 如果没有序号，添加序号
    if not has_number:
        # 清除原有内容
        for run in paragraph.runs:
            run.text = ''
        # 添加序号和原文本
        if paragraph.runs:
            paragraph.runs[0].text = num_text + text
        else:
            paragraph.add_run(num_text + text)


def format_document(input_path, output_path):
    """格式化文档"""
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"错误：无法读取文档 - {e}")
        return False
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(37)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)
    
    # 获取所有段落
    paragraphs = list(doc.paragraphs)
    
    if not paragraphs:
        print("错误：文档没有段落")
        return False
    
    # 检测标题（第一个非空段落）
    title_idx = -1
    for i, p in enumerate(paragraphs):
        if p.text.strip():
            title_idx = i
            break
    
    # 处理每个段落：只设置格式，不修改文本内容
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        
        if not text:
            continue
        
        # 标题处理（第一个非空段落）
        if i == title_idx:
            set_paragraph_format(p, FONT_TITLE, SIZE_TITLE, 
                               WD_ALIGN_PARAGRAPH.CENTER, False, False, is_heading=False)
            continue
        
        # 日期行处理
        if is_date_line(text):
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.CENTER, False, False, is_heading=False)
            continue
        
        # 主送单位处理
        if is_main_subject(text):
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.LEFT, False, False, is_heading=False)
            continue
        
        # 附件说明处理
        if is_attachment(text):
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.LEFT, False, True, is_heading=False)
            continue
        
        # 发文单位处理（右对齐，通常在文档末尾）
        if i > len(paragraphs) - 5 and is_sender(text):
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.RIGHT, False, False, is_heading=False)
            continue
        
        # 检测标题层级（根据已有序号）
        level = detect_heading_level(text)
        
        if level == 1:
            # 一级标题：黑体
            set_paragraph_format(p, FONT_HEADING1, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.JUSTIFY, True, False, is_heading=True)
        elif level == 2:
            # 二级标题：仿宋_GB2312 加粗
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.JUSTIFY, True, True, is_heading=True)
        elif level == 3:
            # 三级标题：仿宋_GB2312 不加粗
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.JUSTIFY, True, False, is_heading=True)
        elif level == 4:
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.JUSTIFY, True, False, is_heading=True)
        # 检测短标题（没有序号的标题，作为一级标题处理，不添加序号）
        elif is_short_heading(text):
            set_paragraph_format(p, FONT_HEADING1, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.JUSTIFY, True, False, is_heading=True)
        else:
            # 默认正文
            set_paragraph_format(p, FONT_BODY, SIZE_BODY, 
                               WD_ALIGN_PARAGRAPH.JUSTIFY, True, False, is_heading=False)
    
    # 保存文档
    try:
        doc.save(output_path)
        print(f"成功：文档已保存至 {output_path}")
        return True
    except Exception as e:
        print(f"错误：无法保存文档 - {e}")
        return False


def main():
    if len(sys.argv) != 3:
        print("用法：python format_gxgk.py <input.docx> <output.docx>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    success = format_document(input_path, output_path)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
